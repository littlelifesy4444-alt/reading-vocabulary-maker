
import os, io, json, base64, random, tempfile, re
from pathlib import Path
from flask import Flask, render_template, request, jsonify, send_file, session
from openai import OpenAI
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import fitz

app = Flask(__name__)
app.secret_key = os.environ.get("FLASK_SECRET_KEY", "change-me-in-production")
client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
MODEL = os.environ.get("OPENAI_MODEL", "gpt-5.6")

SYSTEM_RULES = """
You are an English reading-based vocabulary test generator for Korean students.

CORE PURPOSE
- The test must reward students who reviewed the reading passages.
- Do not make a simple vocabulary-list memorization test.
- Select vocabulary flexibly based on passage difficulty and vocabulary density.
- Do not force a fixed number per passage.
- Single words, phrasal verbs, idioms, collocations, and reading-useful expressions may all be selected.
- Prefer words/expressions important for comprehension or reusable in academic reading.
- Do not select trivial proper nouns unless pedagogically necessary.

VOCABULARY LIST
- Fields: word/expression, pronunciation, POS, Korean meaning, synonym, antonym, new example.
- Synonyms/antonyms must match the target sense used in context.
- If no accurate, context-appropriate synonym or antonym exists, use "—". Never force one.
- New examples must be newly written and MUST NOT be reused in the test.

TEST DESIGN
- Default target for about 5 passages: around 30 questions, but adjust flexibly to input size and real vocabulary value.
- If there are fewer passages, reduce the list and question count rather than forcing filler.
- Prefer about 5–6 tested words per passage overall, but redistribute freely: easy passage may contribute few, difficult passage may contribute many.
- One tested target per question by default; avoid repeating the same target unless necessary.
- Section 1: 40%, subjective, first-letter cue. Reading-review type.
  * Base each item on the original passage.
  * Preserve at least one core fact, situation, causal relation, or contrast.
  * Modify the original sentence only moderately.
  * Do not make syntax harder than the source passage.
  * A student who reviewed the passage should recognize the situation.
  * Do not merely copy the source sentence and blank one word.
- Section 2: 30%, subjective, first-letter cue. New-context transfer type.
  * Use a completely new context.
  * Never reuse the vocabulary-list example.
- Section 3: 15%, mainly multiple choice. English-definition type.
- Section 4: 15%, mainly multiple choice. Synonym/antonym/lexical-relation type.
  * Only use targets for which the relation is genuinely accurate.
- Round section counts to integers while preserving the proportions as closely as possible.
- The order of tested vocabulary must NOT follow the vocabulary-list order.
- Shuffle target order within each section.
- Avoid placing vocabulary that was adjacent in the list next to each other on the test when possible.

DIFFICULTY
- Infer the approximate reading level from the uploaded text itself.
- Difficulty setting:
  easier = more contextual clues, simpler distractors/definitions, source-close Section 1.
  normal = balanced.
  harder = fewer clues, more plausible distractors/definitions, but never exceed the passage's grammatical level excessively.

ANSWER/EXPLANATION
- Student-facing answer sheet: number, answer, Korean meaning, concise explanation.
- Explanations should explain why the word fits the sentence or the lexical relationship.
- Do not label question types in the answer sheet.

ACCURACY
- Never invent passage facts.
- If the source is unreadable or insufficient, say so in source_note and reduce output rather than hallucinating.
"""

SCHEMA = {
  "type": "object",
  "additionalProperties": False,
  "properties": {
    "title": {"type":"string"},
    "detected_level": {"type":"string"},
    "source_note": {"type":"string"},
    "vocabulary": {
      "type":"array",
      "items":{
        "type":"object","additionalProperties":False,
        "properties":{
          "id":{"type":"integer"},
          "word":{"type":"string"},
          "pronunciation":{"type":"string"},
          "pos":{"type":"string"},
          "meaning_ko":{"type":"string"},
          "synonym":{"type":"string"},
          "antonym":{"type":"string"},
          "example":{"type":"string"},
          "passage_label":{"type":"string"},
          "importance":{"type":"integer","minimum":1,"maximum":5}
        },
        "required":["id","word","pronunciation","pos","meaning_ko","synonym","antonym","example","passage_label","importance"]
      }
    },
    "questions":{
      "type":"array",
      "items":{
        "type":"object","additionalProperties":False,
        "properties":{
          "number":{"type":"integer"},
          "section":{"type":"integer","enum":[1,2,3,4]},
          "target":{"type":"string"},
          "prompt":{"type":"string"},
          "choices":{"type":"array","items":{"type":"string"}},
          "answer":{"type":"string"},
          "meaning_ko":{"type":"string"},
          "explanation_ko":{"type":"string"}
        },
        "required":["number","section","target","prompt","choices","answer","meaning_ko","explanation_ko"]
      }
    }
  },
  "required":["title","detected_level","source_note","vocabulary","questions"]
}

def data_url_from_bytes(data, mime):
    return f"data:{mime};base64," + base64.b64encode(data).decode("ascii")

def extract_pdf_pages(data: bytes, page_range: str|None) -> bytes:
    if not page_range or not page_range.strip():
        return data
    src=fitz.open(stream=data, filetype="pdf")
    out=fitz.open()
    wanted=set()
    for part in page_range.split(","):
        part=part.strip()
        if not part: continue
        if "-" in part:
            a,b=part.split("-",1)
            a=int(a.strip()); b=int(b.strip())
            for p in range(min(a,b), max(a,b)+1): wanted.add(p)
        else:
            wanted.add(int(part))
    valid=[p for p in sorted(wanted) if 1 <= p <= len(src)]
    if not valid:
        raise ValueError("입력한 페이지 범위가 PDF 페이지 범위와 맞지 않습니다.")
    for p in valid:
        out.insert_pdf(src, from_page=p-1, to_page=p-1)
    result=out.tobytes()
    src.close(); out.close()
    return result

def build_input(files, page_range):
    content=[]
    for f in files:
        raw=f.read()
        name=(f.filename or "").lower()
        mime=f.mimetype or ""
        if name.endswith(".pdf") or mime=="application/pdf":
            raw=extract_pdf_pages(raw, page_range)
            content.append({
                "type":"input_file",
                "filename":f.filename or "reading.pdf",
                "file_data":data_url_from_bytes(raw,"application/pdf"),
                "detail":"high"
            })
        elif mime.startswith("image/") or name.endswith((".jpg",".jpeg",".png",".webp")):
            mime=mime if mime.startswith("image/") else "image/jpeg"
            content.append({"type":"input_image","image_url":data_url_from_bytes(raw,mime)})
        else:
            raise ValueError("현재 1차 버전은 PDF/JPG/PNG/WEBP 파일을 지원합니다.")
    return content

def generate(title, difficulty, files, page_range):
    if not os.environ.get("OPENAI_API_KEY"):
        raise RuntimeError("서버에 OPENAI_API_KEY가 설정되어 있지 않습니다.")
    content=build_input(files,page_range)
    user_text=f"""
시험지 제목: {title or 'Reading Vocabulary Review'}
난이도 조절: {difficulty}
업로드한 자료의 실제 독해 지문을 분석하여 단어장, 시험지, 학생용 정답·해설지를 한 세트로 생성하세요.
페이지 범위가 적용되어 있다면 제공된 파일의 해당 범위만 사용하세요.
"""
    content.append({"type":"input_text","text":user_text})
    resp=client.responses.create(
        model=MODEL,
        instructions=SYSTEM_RULES,
        input=[{"role":"user","content":content}],
        text={
          "format":{
            "type":"json_schema",
            "name":"reading_vocab_package",
            "strict":True,
            "schema":SCHEMA
          }
        }
    )
    data=json.loads(resp.output_text)
    # deterministic app-side shuffle within section, then renumber.
    by_sec={1:[],2:[],3:[],4:[]}
    for q in data["questions"]:
        by_sec[q["section"]].append(q)
    rng=random.SystemRandom()
    shuffled=[]
    for sec in (1,2,3,4):
        rng.shuffle(by_sec[sec])
        shuffled.extend(by_sec[sec])
    for i,q in enumerate(shuffled,1): q["number"]=i
    data["questions"]=shuffled
    return data

def set_page_border(section):
    sectPr=section._sectPr
    pg=OxmlElement("w:pgBorders")
    pg.set(qn("w:offsetFrom"),"page")
    for edge in ("top","left","bottom","right"):
        el=OxmlElement(f"w:{edge}")
        el.set(qn("w:val"),"single"); el.set(qn("w:sz"),"8")
        el.set(qn("w:space"),"12"); el.set(qn("w:color"),"666666")
        pg.append(el)
    sectPr.append(pg)

def add_footer_motto(section):
    section.footer_distance=Inches(.12)
    p=section.footer.paragraphs[0]
    p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    r=p.add_run("Hard work pays off.")
    r.italic=True; r.font.name="Arial"; r.font.size=Pt(9.5)

def doc_base():
    d=Document()
    s=d.sections[0]
    s.top_margin=Inches(.55); s.bottom_margin=Inches(.55)
    s.left_margin=Inches(.48); s.right_margin=Inches(.48)
    set_page_border(s); add_footer_motto(s)
    d.styles["Normal"].font.name="Arial"
    d.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"),"Malgun Gothic")
    return d

def heading(d,title,subtitle):
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(title); r.bold=True; r.font.size=Pt(18)
    p.paragraph_format.space_after=Pt(0)
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(subtitle); r.bold=True; r.font.size=Pt(12)
    p.paragraph_format.space_after=Pt(6)

def save_vocab_docx(data,path):
    d=doc_base(); heading(d,data["title"],"Vocabulary List")
    widths=[.32,.80,1.02,.92,1.08,3.94]
    t=d.add_table(rows=1,cols=6); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    headers=["No.","Word","발음·품사","뜻","Syn./Ant.","New Example"]
    for i,h in enumerate(headers):
        t.rows[0].cells[i].text=h
        for r in t.rows[0].cells[i].paragraphs[0].runs: r.bold=True; r.font.size=Pt(10)
    for idx,v in enumerate(data["vocabulary"],1):
        rel=[]
        if v["synonym"]!="—": rel.append("S: "+v["synonym"])
        if v["antonym"]!="—": rel.append("A: "+v["antonym"])
        vals=[str(idx),v["word"],f'{v["pronunciation"]}\n{v["pos"]}',v["meaning_ko"],"\n".join(rel) if rel else "—",v["example"]]
        cells=t.add_row().cells
        for i,x in enumerate(vals):
            cells[i].text=x
            cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.0
                for r in p.runs: r.font.size=Pt(10.4 if i in (1,5) else 9.8)
    p=d.add_paragraph("※ 예문은 단어 학습용 새 문장이며 시험에는 재사용하지 않습니다. 정확한 동의어·반의어가 없으면 —로 표시합니다.")
    for r in p.runs: r.font.size=Pt(9)
    d.save(path)

def save_test_docx(data,path):
    d=doc_base(); heading(d,data["title"],"Vocabulary Review Test")
    info=d.add_table(rows=1,cols=3)
    for i,x in enumerate(["학원/반: ____________________","이름: ____________________","점수: ______ / ______"]):
        info.cell(0,i).text=x
        for r in info.cell(0,i).paragraphs[0].runs: r.bold=True; r.font.size=Pt(11.5)
    p=d.add_paragraph("1·2: 첫 글자를 참고하여 알맞은 단어를 쓰시오.   3·4: 가장 알맞은 답을 고르시오.")
    for r in p.runs: r.bold=True; r.font.size=Pt(11.5)
    qs=data["questions"]
    # 2-column, roughly half per page. Keep large spacing.
    per_page=16
    for page_start in range(0,len(qs),per_page):
        if page_start>0: d.add_page_break()
        batch=qs[page_start:page_start+per_page]
        split=(len(batch)+1)//2
        tt=d.add_table(rows=1,cols=2); tt.autofit=False
        for ci,sub in enumerate([batch[:split],batch[split:]]):
            cell=tt.cell(0,ci)
            seen=set()
            for q in sub:
                if q["section"] not in seen:
                    seen.add(q["section"])
                    p=cell.add_paragraph()
                    r=p.add_run(str(q["section"])); r.bold=True; r.font.size=Pt(14)
                p=cell.add_paragraph()
                p.paragraph_format.space_after=Pt(14); p.paragraph_format.line_spacing=1.35
                text=f'{q["number"]}. {q["prompt"]}'
                r=p.add_run(text); r.font.size=Pt(12.5)
                if q["choices"]:
                    p.add_run("\n")
                    r=p.add_run("   ".join(q["choices"])); r.font.size=Pt(11.8)
    d.save(path)

def save_answer_docx(data,path):
    d=doc_base(); heading(d,data["title"],"Answer & Explanation")
    t=d.add_table(rows=1,cols=4); t.autofit=False
    headers=["No.","정답","뜻","문제 해설"]
    for i,h in enumerate(headers):
        t.rows[0].cells[i].text=h
        for r in t.rows[0].cells[i].paragraphs[0].runs: r.bold=True; r.font.size=Pt(10.5)
    for q in data["questions"]:
        vals=[str(q["number"]),q["answer"],q["meaning_ko"],q["explanation_ko"]]
        cells=t.add_row().cells
        for i,x in enumerate(vals):
            cells[i].text=x
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.0
                for r in p.runs: r.font.size=Pt(10.5 if i<3 else 10.8)
    d.save(path)

@app.route("/")
def index():
    return render_template("index.html")

@app.post("/api/generate")
def api_generate():
    try:
        files=request.files.getlist("files")
        if not files: return jsonify({"error":"파일을 하나 이상 선택하세요."}),400
        data=generate(
            request.form.get("title","").strip(),
            request.form.get("difficulty","normal"),
            files,
            request.form.get("page_range","").strip()
        )
        token=os.urandom(8).hex()
        session["result_token"]=token
        session["result_json"]=json.dumps(data,ensure_ascii=False)
        return jsonify(data)
    except Exception as e:
        return jsonify({"error":str(e)}),500

@app.post("/api/reshuffle")
def reshuffle():
    raw=session.get("result_json")
    if not raw: return jsonify({"error":"먼저 시험을 생성하세요."}),400
    data=json.loads(raw)
    by={1:[],2:[],3:[],4:[]}
    for q in data["questions"]: by[q["section"]].append(q)
    rng=random.SystemRandom(); qs=[]
    for s in (1,2,3,4):
        rng.shuffle(by[s]); qs += by[s]
    for i,q in enumerate(qs,1): q["number"]=i
    data["questions"]=qs
    session["result_json"]=json.dumps(data,ensure_ascii=False)
    return jsonify(data)

@app.post("/api/regenerate-question")
def regenerate_question():
    raw=session.get("result_json")
    if not raw: return jsonify({"error":"먼저 시험을 생성하세요."}),400
    data=json.loads(raw)
    body=request.get_json(force=True)
    number=int(body["number"]); difficulty=body.get("difficulty","normal")
    q=next((x for x in data["questions"] if x["number"]==number),None)
    if not q: return jsonify({"error":"문항을 찾을 수 없습니다."}),404
    prompt=f"""
기존 문항을 같은 section과 같은 target word로 다시 작성하세요.
난이도: {difficulty}
기존 문항: {json.dumps(q,ensure_ascii=False)}
핵심 규칙은 전체 앱 규칙을 그대로 따르세요.
Section 1이면 원 지문 맥락을 유지해야 하므로, 현재 문항에 포함된 원문 사실을 유지한 채 표현만 조정하세요.
Section 2면 완전히 새 문맥을 사용하세요.
정답과 뜻, 학생용 해설도 함께 반환하세요.
"""
    qschema={"type":"object","additionalProperties":False,"properties":{
        "number":{"type":"integer"},"section":{"type":"integer","enum":[1,2,3,4]},
        "target":{"type":"string"},"prompt":{"type":"string"},
        "choices":{"type":"array","items":{"type":"string"}},"answer":{"type":"string"},
        "meaning_ko":{"type":"string"},"explanation_ko":{"type":"string"}},
        "required":["number","section","target","prompt","choices","answer","meaning_ko","explanation_ko"]}
    resp=client.responses.create(
        model=MODEL, instructions=SYSTEM_RULES,
        input=prompt,
        text={"format":{"type":"json_schema","name":"question","strict":True,"schema":qschema}}
    )
    newq=json.loads(resp.output_text); newq["number"]=number; newq["section"]=q["section"]; newq["target"]=q["target"]
    data["questions"]=[newq if x["number"]==number else x for x in data["questions"]]
    session["result_json"]=json.dumps(data,ensure_ascii=False)
    return jsonify(newq)

@app.get("/download/<kind>")
def download(kind):
    raw=session.get("result_json")
    if not raw: return "먼저 시험을 생성하세요.",400
    data=json.loads(raw)
    tmp=tempfile.NamedTemporaryFile(delete=False,suffix=".docx")
    tmp.close()
    if kind=="vocab":
        save_vocab_docx(data,tmp.name); name="Vocabulary_List.docx"
    elif kind=="test":
        save_test_docx(data,tmp.name); name="Vocabulary_Test.docx"
    elif kind=="answer":
        save_answer_docx(data,tmp.name); name="Answer_Explanation.docx"
    else:
        return "잘못된 요청",404
    return send_file(tmp.name,as_attachment=True,download_name=name)

if __name__ == "__main__":
    app.run(host="0.0.0.0",port=int(os.environ.get("PORT","5000")),debug=True)
